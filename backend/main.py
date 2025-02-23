import shutil
from fastapi import FastAPI, HTTPException, Depends, File, UploadFile
from pymongo import MongoClient
from pydantic import BaseModel
from typing import List, Optional
from jose import jwt, JWTError
from passlib.context import CryptContext
import uuid
import os
import smtplib
import datetime
import pandas as pd
from dotenv import load_dotenv
from keybert import KeyBERT
from fastapi.responses import JSONResponse
import gridfs
from bson import ObjectId
from pdfminer.high_level import extract_text

from docx import Document as DocxDocument

# Load environment variables
load_dotenv()

# MongoDB Atlas Connection
MONGO_URI = os.getenv("MONGO_URI")
client = MongoClient(MONGO_URI)
db = client["inventory_docs"]
fs = gridfs.GridFS(db)

doc_collection = db["documents"]
inv_collection = db["inventory"]
user_collection = db["users"]
log_collection = db["audit_logs"]

kw_model = KeyBERT()
# FastAPI Setup
app = FastAPI()

# Security Setup
SECRET_KEY = os.getenv("SECRET_KEY")
ALGORITHM = "HS256"
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# Email Setup
EMAIL_HOST = os.getenv("EMAIL_HOST")
EMAIL_PORT = os.getenv("EMAIL_PORT")
EMAIL_USERNAME = os.getenv("EMAIL_USERNAME")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")

# Models
class User(BaseModel):
    username: str
    password: str
    role: str  # admin, manager, viewer

class Document(BaseModel):
    title: str
    content: str
    template: Optional[str] = "default"
    tags: List[str] = []  # Ensure this field is included

class LoginRequest(BaseModel):
    username: str
    password: str
# Extract text from files
def extract_text_from_file(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File {file_path} not found!")

    if file_path.endswith(".pdf"):
        return extract_text(file_path)  # Extract text from PDF
    elif file_path.endswith(".docx"):
        doc = DocxDocument(file_path)
        return "\n".join([para.text for para in doc.paragraphs])  # Extract text from Word file
    else:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()  # Extract text from text files

# Authentication Helpers
def get_hashed_password(password):
    return pwd_context.hash(password)

def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def generate_token(username: str, role: str):
    return jwt.encode({"sub": username, "role": role}, SECRET_KEY, algorithm=ALGORITHM)

def authenticate_user(username: str, password: str):
    user = user_collection.find_one({"username": username})
    if user and verify_password(password, user["password"]):
        return user
    return None

# Free AI-Driven Document Tagging
def tag_document(text):
    keywords = kw_model.extract_keywords(text, keyphrase_ngram_range=(1, 2), stop_words='english', top_n=5)
    return [kw[0] for kw in keywords]

# Centralized Documentation Repository
def store_document_with_template(doc: Document):
    doc_id = str(uuid.uuid4())
    doc.tags = tag_document(doc.content)
    doc_collection.insert_one({"_id": doc_id, "title": doc.title, "content": doc.content, "tags": doc.tags, "template": doc.template})
    return {"message": "Document created with structured template", "id": doc_id, "tags": doc.tags}


# API Endpoints
@app.post("/register/")
def register_user(user: User):
    if user_collection.find_one({"username": user.username}):
        raise HTTPException(status_code=400, detail="User already exists")
    hashed_password = get_hashed_password(user.password)
    user_collection.insert_one({"username": user.username, "password": hashed_password, "role": user.role})
    return {"message": "User registered successfully"}

@app.post("/login/")
def login_user(login_request: LoginRequest):
    user_data = login_request.dict()  # ✅ Ensures correct data format
    authenticated_user = authenticate_user(user_data["username"], user_data["password"])

    if not authenticated_user:
        raise HTTPException(status_code=400, detail="Invalid credentials")

    token = generate_token(authenticated_user["username"], authenticated_user["role"])
    return {"access_token": token, "role": authenticated_user["role"]}

@app.post("/documents/")
def create_document(doc: Document):
    return store_document_with_template(doc)
def extract_text_from_file(file_path):
    if file_path.endswith(".pdf"):
        return extract_text(file_path)  
    elif file_path.endswith(".docx"):
        doc = DocxDocument(file_path)
        return "\n".join([para.text for para in doc.paragraphs])  
    else:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    file_location = os.path.join('./', file.filename)

    # Save file to disk
    with open(file_location, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # Extract text AFTER saving the file
    extracted_text = extract_text_from_file(file_location)

    # Generate tags using AI tagging method
    tags = tag_document(extracted_text)

    # Store document metadata in MongoDB
    doc_id = str(uuid.uuid4())
    db.documents.insert_one({
        "_id": doc_id,
        "filename": file.filename,
        "extracted_text": extracted_text,
        "tags": tags
    })

    return {"message": "File uploaded successfully", "file_id": doc_id, "content": extracted_text, "tags": tags}

# Retrieve File from GridFS
@app.get("/download/{file_id}")
def download_file(file_id: str):
    file = fs.get(ObjectId(file_id))  # Retrieve file
    return {"filename": file.filename, "content": file.read()}  # Returns file data

# Retrieve Extracted Text
@app.get("/documents/")
def list_documents():
    return list(db.documents.find({}, {"_id": 0}))  # List all uploaded documents
@app.get("/documents/{doc_id}")
def get_document(doc_id: str):
    doc = doc_collection.find_one({"_id": doc_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc

@app.get("/documents/")
def list_documents():
    return list(doc_collection.find({}, {"_id": 0}))

@app.put("/documents/{doc_id}")
def update_document(doc_id: str, doc: Document):
    doc.tags = tag_document(doc.content)
    updated_doc = {"title": doc.title, "content": doc.content, "template": doc.template, "tags": doc.tags}
    result = doc_collection.update_one({"_id": doc_id}, {"$set": updated_doc})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"message": "Document updated successfully", "tags": doc.tags}


